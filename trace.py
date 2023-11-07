import os
import pandas as pd
from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from docx.shared import Pt


def extract_requirements_from_word_docs():
    style1 = ""
    style2 = ""
    data = []
    directory_path = 'd:\\work\\tmp\\'
    if not os.path.exists(directory_path):
        print(f"Directory '{directory_path}' does not exist.")
        return

    for foldername, subfolders, filenames in os.walk(directory_path):
        for filename in filenames:
            if filename.endswith('.docx'):
                full_path = os.path.join(foldername, filename)
                try:
                    document = Document(full_path)
                    for table in document.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                if 'Requirement_ID' in [p.style.name for p in cell.paragraphs]:
                                    style1 = cell.text
                                if 'Requirement_Text' in [p.style.name for p in cell.paragraphs]:
                                    style2 = cell.text
                        if len(style1) == 0 and len(style2) == 0:
                            continue
                        data.append([style1, style2])
                        style1 = ""
                        style2 = ""
                except PackageNotFoundError:
                    print(f'Could not open file: {full_path}')  # Отображем полный путь

    df = pd.DataFrame(data, columns=['style1', 'style2'])
    df.to_excel('d:\\work\\tmp\\output.xlsx', index=False)

extract_requirements_from_word_docs()